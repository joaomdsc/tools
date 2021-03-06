#!/usr/bin/env python
# docx_common.py - common functions for the creation of  Word .docx files

import re

# Writing out to Word .docx files
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_BREAK

# Colors
black = RGBColor(0x0, 0x0, 0x0)
char_red = RGBColor(0xd7, 0x56, 0x45)     # (215, 86,  69) dec
pinyin_blue = RGBColor(0x2f, 0x54, 0x96)  # ( 47, 84, 150) dec
      
#--------------------------------------------------------------------------
# Add page break
#--------------------------------------------------------------------------

def add_page_break(doc):
    # One paragraph for everyting
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)

#--------------------------------------------------------------------------
# Add one character run
#--------------------------------------------------------------------------

def add_char_run(p, 汉字, font='KaiTi', color=char_red, strike=None):
    # One run for a chinese character
    r = p.add_run()
    
    f = r.font
    f.name = 'KaiTi'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    f.size = Pt(16)
    f.color.rgb = color
    if strike:
        f.strike = strike

    r.text = 汉字

#--------------------------------------------------------------------------
# Add one pinyin run
#--------------------------------------------------------------------------

def add_pinyin_run(p, pinyin, color=pinyin_blue, strike=None):
    # One run for the pinyin
    r = p.add_run()
    
    f = r.font
    f.name = 'Calibri'
    f.size = Pt(10)
    f.color.rgb = color
    if strike:
        f.strike = strike

    r.text = pinyin
  
#--------------------------------------------------------------------------
# Add a text run
#--------------------------------------------------------------------------

def add_text_run(p, text, bold=None, italic=None, strike=None, bg=None):
    # One run for plain text
    r = p.add_run()
    
    f = r.font
    # f.name = 'Times New Roman'
    f.name = 'Calibri'
    f.size = Pt(10)
    # if bg:
    #     f.highlight_color = bg

    r.text = text
    if bold:
        r.bold = bold
    if italic:
        r.italic = italic
  
#--------------------------------------------------------------------------
# Add an HTML run (plain text with interspersed HTML)
#--------------------------------------------------------------------------

def add_html_run(p, s, bg=None):
    # One run for text with optional embedded HTML
    cc = 0
    pat = re.compile('<(i|span) class="(hanzi|pinyin)">([^<]+)</(i|span)>')
    m = re.search(pat, s)
    while m:
        add_text_run(p, s[:m.start()])
        if m.group(2) == 'hanzi':
            add_char_run(p, m.group(3))
        else:
            add_pinyin_run(p, m.group(3))
        s = s[m.end():]
        m = re.search(pat, s)
    add_text_run(p, s, bg=bg)
    
#--------------------------------------------------------------------------
# Create a new paragraph object
#--------------------------------------------------------------------------

def new_paragraph(doc, style=None):
    p = doc.add_paragraph(style=style)

    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    return p

#--------------------------------------------------------------------------
# Write a paragraph with simple text content
#--------------------------------------------------------------------------

def add_text_paragraph(doc, text, bg=None, style=None):
    p = new_paragraph(doc, style=style)
    add_text_run(p, text, bg=bg)
    
#--------------------------------------------------------------------------
# Write a paragraph with HTML text content
#--------------------------------------------------------------------------

def add_paragraph(doc, text, bg=None):
    p = new_paragraph(doc)
    add_html_run(p, text)

#--------------------------------------------------------------------------
# Write one heading with style Heading 1, 2, or 3
#--------------------------------------------------------------------------

def add_heading(doc, level, text):
    # One paragraph for the heading line
    p = doc.add_paragraph()
    p.style = doc.styles[f'Heading {level}']
    r = p.add_run()

    # Writing hanzi characters in the same font (but bold and italic are not
    # very legible)
    pat = re.compile('<i class="hanzi">([^<]+)</i>')
    m = re.search(pat, text)
    r.text = text[:m.start()] + m.group(1) + text[m.end():] if m else text
    
# -----------------------------------------------------------------------------
# main
# -----------------------------------------------------------------------------
    
if __name__ == '__main__':
    print('This module is not meant to be executed directly')